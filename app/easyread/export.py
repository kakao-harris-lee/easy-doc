"""검수 완료 문서의 내보내기 산출물 생성 (docx·txt·hwpx).

**hwpx 한계**: OWPML 패키지는 스펙과 우리 ingest 파서가 읽는 구조를 보고 직접 조립한다
(app/easyread/hwpx.py). 자동으로 확인하는 것은 "우리 추출기로 다시 읽힌다"까지이고,
**한컴 오피스에서 실제로 열리는지는 파일럿에서 사람이 확인해야 한다** — 저장소·CI에
한컴 오피스가 없다.

**이 모듈의 `restore_placeholders`는 내보내기 경로 전용이다.** 마스킹은 문서 본문이
LLM으로 나가는 것을 막는 보안 불변식(master-plan 3.2)이고, 내보내기는 그 불변식의
유일한 예외다 — 담당자가 그대로 배포할 최종 문서에 `[[주민등록번호1]]`이 남아 있으면
쓸 수 없기 때문이다. 조회 응답·목록·로그 등 다른 경로에서 이 함수를 부르지 않는다.
원문 개인정보가 흘러 다니는 표면을 한 곳으로 묶어두는 것이 이 규칙의 목적이다.

산출물에는 "AI 초안" 같은 꼬리말을 붙이지 않는다 — 담당자가 지우고 배포해야 하는
문구를 우리가 만들어 넣을 이유가 없다. AI 초안임을 알리는 책임은 검수 화면(HITL
배너, master-plan 3.3)에 있다.

제어문자 정규화는 `render_export` 진입부에서 한 번 한다. 본문 재료 중 AI 초안과 마스킹
원문은 저장 경로에서 정규화되지 않으므로 여기가 **유일한** 방어다(제목·검수 수정본은
저장 시점에 이미 다듬어져 들어온다 — 사유는 app/text.py).
"""

import io
import re
from collections.abc import Mapping
from dataclasses import dataclass
from enum import StrEnum
from pathlib import PurePosixPath
from typing import assert_never
from urllib.parse import quote

import docx

from app.easyread.hwpx import MIMETYPE as HWPX_MIMETYPE
from app.easyread.hwpx import build_hwpx
from app.text import strip_control_chars


class ExportFormat(StrEnum):
    """내보내기 파일 형식. 값이 그대로 확장자이자 쿼리 파라미터 값이다.

    pdf와 구버전 hwp(바이너리)는 Lean MVP 범위 밖이다. 여기에 이름이 없는 형식은
    라우터의 쿼리 검증에서 422로 걸린다.
    """

    DOCX = "docx"
    TXT = "txt"
    HWPX = "hwpx"


#: 형식별 미디어 타입. 라우터가 OpenAPI 문서에도 그대로 싣는다 — 문자열을 두 곳에
#: 적어두면 형식이 늘 때 문서만 낡는다.
MEDIA_TYPES: Mapping[ExportFormat, str] = {
    ExportFormat.DOCX: "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    # charset 명시: txt는 BOM 없이 UTF-8로 내보내므로, 이 표시가 없으면 브라우저가
    # 로캘 기본 인코딩으로 열어 한글이 깨진다.
    ExportFormat.TXT: "text/plain; charset=utf-8",
    # hwpx에는 IANA 등록 미디어 타입이 없다. 후보 둘 중 `application/hwp+zip`을 쓰는
    # 이유는 근거가 세 곳에서 일치하기 때문이다: 한컴 개발자 포럼의 공식 답변("공식
    # 등록된 mimetype이 없으니 커스텀 동작을 구현한다면 application/hwp+zip"),
    # Apache Tika의 형식 정의, 그리고 우리가 만드는 패키지 안 mimetype 항목의 값
    # (OWPML 규약). `application/vnd.hancom.hwpx`는 한컴 리눅스 패키지의 데스크톱
    # 등록값으로 도는 표기지만 역시 미등록이고, 패키지 내부 표기와 어긋난다.
    ExportFormat.HWPX: HWPX_MIMETYPE,
}

#: 파일명에서 걷어낼 문자 — 경로 구분자, 제어문자, 따옴표. 제목은 본문 첫 줄에서
#: 유도한 사용자 입력이라 무엇이든 들어올 수 있다.
_FORBIDDEN_IN_FILENAME = re.compile(r'[\x00-\x1f\x7f"\\/:*?<>|]')

#: 제목이 전부 걸러졌을 때 쓰는 이름.
_FALLBACK_NAME = "쉬운 글"

#: 파일명 본체 최대 길이(문자). 파일 시스템은 대개 255**바이트**가 한계인데 한글은
#: 글자당 3바이트라, 확장자까지 더해도 넉넉히 들어가는 길이로 자른다.
_MAX_FILENAME_STEM = 80

#: ASCII 대체 파일명의 본체. 헤더는 latin-1만 담을 수 있어 한글 이름을 여기 넣을 수
#: 없다 — 진짜 이름은 filename*(UTF-8)으로 보내고 이쪽은 옛 클라이언트용 자리다.
_ASCII_FALLBACK_STEM = "easy-read"

#: 문단 경계. 빈 줄(공백만 있는 줄 포함)이 하나 이상이면 새 문단으로 본다.
_PARAGRAPH_BREAK = re.compile(r"\n\s*\n")

#: 자리표시자 한 개. 대괄호가 포함되지 않은 라벨만 받는다(`[[주민등록번호1]]`).
_PLACEHOLDER = re.compile(r"\[\[[^\[\]]+\]\]")


@dataclass(frozen=True)
class ExportFile:
    """내보낼 파일 한 건. 전송에 필요한 것만 담는다(HTTP 헤더 조립은 라우터가 한다)."""

    filename: str
    media_type: str
    content: bytes


def restore_placeholders(text: str, originals: Mapping[str, str]) -> str:
    """자리표시자를 원문 값으로 되돌린다 (**내보내기 전용** — 모듈 docstring 참고).

    Args:
        text: 자리표시자가 남아 있는 변환 결과.
        originals: 자리표시자 → 가려졌던 원문 값.

    자리표시자를 하나씩 찾아 한 번에 바꾼다(단일 패스). 치환 결과를 다시 훑지 않으므로
    복원된 원문이 우연히 자리표시자 모양이어도 두 번 치환되지 않고, 항목 수만큼 본문
    전체를 다시 쓰지도 않는다.

    목록에 없는 자리표시자는 그대로 둔다 — 우리가 만든 것이 아닌 표기를 지워
    본문을 조용히 훼손하지 않는다.
    """
    return _PLACEHOLDER.sub(lambda match: originals.get(match.group(), match.group()), text)


def export_filename(title: str, export_format: ExportFormat) -> str:
    """문서 제목으로 내려받을 파일명을 만든다.

    제목은 본문 첫 줄에서 유도한 사용자 입력이므로 경로 구분자·제어문자를 걷어낸다 —
    파일명이 디렉터리를 벗어나거나(../) 저장이 실패하는 이름이 되지 않게 한다.
    """
    cleaned = " ".join(_FORBIDDEN_IN_FILENAME.sub(" ", title).split())
    # 앞뒤 점 제거: 숨김 파일(.name)이나 윈도우가 거부하는 이름(name.)이 되는 것을 막는다.
    stem = cleaned.strip(". ")[:_MAX_FILENAME_STEM].strip(". ")
    return f"{stem or _FALLBACK_NAME}.{export_format.value}"


def content_disposition(filename: str) -> str:
    """RFC 5987 방식으로 파일명을 담은 Content-Disposition 헤더 값을 만든다.

    HTTP 헤더 값은 latin-1만 담을 수 있어 한글 파일명을 그대로 넣으면 인코딩 오류이거나
    깨진 이름이 된다. 퍼센트 인코딩한 `filename*`을 진짜 이름으로 쓰고, 그 확장을 모르는
    클라이언트를 위해 ASCII 대체 이름을 함께 둔다.
    """
    suffix = PurePosixPath(filename).suffix
    return (
        f'attachment; filename="{_ASCII_FALLBACK_STEM}{suffix}"; '
        f"filename*=UTF-8''{quote(filename, safe='')}"
    )


def render_export(*, export_format: ExportFormat, title: str, body: str) -> ExportFile:
    """최종 본문을 파일 바이트로 만든다.

    Args:
        export_format: 산출물 형식.
        title: 문서 제목 (docx의 제목 문단과 파일명에 쓰인다).
        body: 내보낼 본문 — 자리표시자는 이미 복원된 상태여야 한다.

    제어문자 정규화를 **이 진입부 한 곳**에서 한다. 본문을 이루는 두 재료(워커가 저장한
    AI 초안, 복원해 넣은 마스킹 원문)는 저장 경로에서 정규화되지 않으므로 여기가 유일한
    방어다(모듈 docstring 참고). 형식별 분기 뒤가 아니라 앞에서 거르는 이유는 docx와 txt가
    같은 본문을 내놓아야 하기 때문이다 — XML만 못 담는 문자라고 docx에서만 지우면 두
    산출물의 내용이 갈린다.
    """
    title = strip_control_chars(title)
    body = strip_control_chars(body)
    match export_format:
        case ExportFormat.DOCX:
            content = _render_docx(title=title, body=body)
        case ExportFormat.TXT:
            # txt에는 제목 줄을 덧붙이지 않는다. docx는 제목 스타일이 구조를 만들지만
            # txt는 구조가 없어, 제목을 얹으면 본문과 구분되지 않는 중복 줄만 남는다.
            content = body.encode("utf-8")
        case ExportFormat.HWPX:
            content = _render_hwpx(title=title, body=body)
        case _:
            # 형식을 추가하고 여기를 잊으면 mypy가 이 줄에서 잡는다(런타임까지 가지 않는다).
            assert_never(export_format)
    return ExportFile(
        filename=export_filename(title, export_format),
        media_type=MEDIA_TYPES[export_format],
        content=content,
    )


def _render_docx(*, title: str, body: str) -> bytes:
    """제목 + 본문 문단으로 docx 파일을 만든다.

    인자는 이미 정규화된 상태로 들어온다(`render_export`) — 여기서 다시 거르지 않는 것은
    "어디서 걸렀나"를 찾아 두 곳을 뒤지게 만들지 않기 위해서다. 제어문자가 남은 채
    닿으면 lxml이 ValueError를 던진다.
    """
    document = docx.Document()
    document.add_heading(title, level=1)
    for block in _split_paragraphs(body):
        # python-docx가 문단 안의 \n을 <w:br/>로 옮긴다 — 줄만 바꾼 곳(목록 등)이
        # 한 줄로 뭉치지 않는다.
        document.add_paragraph(block)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _render_hwpx(*, title: str, body: str) -> bytes:
    """제목 + 본문 문단으로 hwpx(OWPML) 패키지를 만든다.

    인자는 이미 정규화된 상태로 들어온다(`render_export`) — docx와 같은 이유다.

    docx는 문단 안의 줄바꿈을 강제 개행(`<w:br/>`)으로 옮기지만 여기서는 **줄마다 문단을
    하나씩** 만든다. 강제 개행(`<hp:lineBreak/>`)까지 조립하면 검증할 수 없는 스키마 조각이
    하나 더 늘어나는데, 목록처럼 줄만 바꾼 곳은 문단으로 나눠도 보이는 결과가 사실상 같다.

    제목은 문단 하나로 얹는다(docx의 제목 문단과 같은 자리). 제목 스타일을 주지 않는 것은
    우리 header.xml이 스타일을 하나만 정의하기 때문이다 — 검증할 수 없는 서식 정의를
    늘리는 대신 본문이 온전히 남는 쪽을 택했다.
    """
    paragraphs: list[str] = []
    if heading := title.strip():
        paragraphs.append(heading)
    for block in _split_paragraphs(body):
        paragraphs.extend(stripped for line in block.splitlines() if (stripped := line.strip()))
    return build_hwpx(paragraphs)


def _split_paragraphs(body: str) -> list[str]:
    """빈 줄을 경계로 본문을 문단 목록으로 나눈다 (빈 문단은 버린다)."""
    return [block.strip() for block in _PARAGRAPH_BREAK.split(body.strip()) if block.strip()]
