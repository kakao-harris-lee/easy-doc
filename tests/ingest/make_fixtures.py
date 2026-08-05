"""테스트용 문서 픽스처 생성 스크립트 (개발자 도구, 테스트가 import 하지 않는다).

생성물은 `tests/ingest/fixtures/`에 커밋한다. 바이너리를 커밋하는 대신 매번
생성하지 않는 이유는 테스트 실행 시간과 재현성 때문이고, 이 스크립트를 함께
커밋하는 이유는 픽스처가 어떤 내용인지·어떻게 다시 만드는지 남기기 위해서다.

재생성::

    uv run python tests/ingest/make_fixtures.py

주의: zip 컨테이너 형식(docx·hwpx)은 항목마다 수정 시각이 들어가 재생성할 때마다
바이트가 달라진다. 내용을 바꾼 게 아니라면 굳이 다시 만들어 커밋하지 않는 편이 좋다
(PDF는 손으로 조립해 시각 정보가 없으므로 재생성해도 동일하다).

픽스처에 담는 텍스트는 전부 가상의 공지문이며 개인정보를 포함하지 않는다.
"""

import zipfile
from collections.abc import Sequence
from pathlib import Path

import docx
from docx.oxml import parse_xml

FIXTURES_DIR = Path(__file__).parent / "fixtures"

# 픽스처에 심는 문장. 테스트가 이 상수를 재선언해 검증하므로 값을 바꾸면 테스트도 고친다.
DOCX_PARAGRAPHS = ("쉬운 글 변환 안내", "이 문서는 추출 테스트용 예시입니다.")
DOCX_TABLE_CELLS = (("구분", "내용"), ("접수 기간", "3월 1일부터 3월 31일까지"))
PDF_PAGES = ("첫째 쪽 안내문입니다.", "둘째 쪽 안내문입니다.")
HWPX_SECTION0 = ("한글 문서 추출 확인", "문단 하나가 여러 조각으로 나뉘어도 이어 붙는다.")
HWPX_SECTION1 = ("둘째 구역의 문장입니다.",)

# sample_rich.docx — python-docx의 paragraphs/tables 순회로는 놓치는 자리들.
RICH_HEAD = "첫 문단입니다."
RICH_TABLE_CELL = "바깥 표 셀"
RICH_NESTED_CELL = "중첩 표 셀"
RICH_TAIL = "표 뒤에 오는 문단입니다."
RICH_TEXTBOX = "텍스트 상자 안 문장입니다."
RICH_INSERTED = "변경 추적으로 삽입된 문장입니다."
RICH_DELETED = "변경 추적으로 삭제된 문장입니다."
RICH_SECOND_SECTION = "둘째 구역 본문입니다."
RICH_HEADER = "머리글 문구"
RICH_FOOTER = "바닥글 문구"

_WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_VML_NS = "urn:schemas-microsoft-com:vml"


def _write_docx_paragraphs() -> None:
    """문단만 있는 docx. 빈 문단을 섞어 공백 정규화까지 확인한다."""
    document = docx.Document()
    document.add_paragraph(DOCX_PARAGRAPHS[0])
    document.add_paragraph("   ")  # 공백뿐인 문단 — 결과에 빈 줄로 남으면 안 된다
    document.add_paragraph(DOCX_PARAGRAPHS[1])
    document.save(str(FIXTURES_DIR / "sample.docx"))


def _write_docx_table() -> None:
    """문단 + 표가 섞인 docx."""
    document = docx.Document()
    document.add_paragraph(DOCX_PARAGRAPHS[0])
    table = document.add_table(rows=len(DOCX_TABLE_CELLS), cols=len(DOCX_TABLE_CELLS[0]))
    for row_index, row_values in enumerate(DOCX_TABLE_CELLS):
        for column_index, value in enumerate(row_values):
            table.cell(row_index, column_index).text = value
    document.save(str(FIXTURES_DIR / "sample_table.docx"))


def _write_docx_rich() -> None:
    """문서 순서 순회로만 잡히는 자리들을 한 파일에 모은 docx.

    - 표 뒤에 문단을 두어, 표가 문서 끝으로 밀리지 않는지 확인한다.
    - 중첩 표·텍스트박스(VML `w:pict`)·변경 추적(`w:ins`/`w:del`)은 python-docx의
      `paragraphs`/`tables` 순회로는 보이지 않는다. 텍스트박스와 변경 추적은
      python-docx API로 만들 수 없어 OOXML 조각을 직접 붙인다.
    - 삭제문은 `w:delText`에 담기므로 `w:t`만 걷는 추출기에서 자연히 빠져야 한다.
    """
    document = docx.Document()
    document.add_paragraph(RICH_HEAD)

    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = RICH_TABLE_CELL
    nested = cell.add_table(rows=1, cols=1)
    nested.cell(0, 0).text = RICH_NESTED_CELL

    document.add_paragraph(RICH_TAIL)

    # VML 텍스트박스. w:txbxContent 안의 문단은 본문 트리에 있지만 Document.paragraphs에는
    # 잡히지 않는다 (run 하위에 들어 있기 때문). 조각은 add_paragraph()가 만든 문단 안에
    # 넣는다 — body에 직접 append하면 끝의 sectPr 뒤로 밀려 형식이 깨진다.
    document.add_paragraph()._p.append(
        parse_xml(
            f'<w:r xmlns:w="{_WORD_NS}" xmlns:v="{_VML_NS}"><w:pict>'
            '<v:shape style="width:200pt;height:50pt"><v:textbox><w:txbxContent>'
            f"<w:p><w:r><w:t>{RICH_TEXTBOX}</w:t></w:r></w:p>"
            "</w:txbxContent></v:textbox></v:shape>"
            "</w:pict></w:r>"
        )
    )

    # 변경 추적. 삽입문은 남고 삭제문은 빠져야 한다.
    tracked = document.add_paragraph()._p
    tracked.append(
        parse_xml(
            f'<w:ins xmlns:w="{_WORD_NS}" w:id="101" w:author="tester"'
            f' w:date="2026-01-01T00:00:00Z"><w:r><w:t>{RICH_INSERTED}</w:t></w:r></w:ins>'
        )
    )
    tracked.append(
        parse_xml(
            f'<w:del xmlns:w="{_WORD_NS}" w:id="102" w:author="tester"'
            f' w:date="2026-01-01T00:00:00Z">'
            f"<w:r><w:delText>{RICH_DELETED}</w:delText></w:r></w:del>"
        )
    )

    section = document.sections[0]
    section.header.paragraphs[0].text = RICH_HEADER
    section.footer.paragraphs[0].text = RICH_FOOTER

    # 둘째 구역. 머리글·바닥글을 앞 구역에서 물려받으므로(is_linked_to_previous),
    # 물려받은 쪽까지 걷는 추출기는 같은 문구를 두 번 내보낸다.
    document.add_section()
    document.add_paragraph(RICH_SECOND_SECTION)

    document.save(str(FIXTURES_DIR / "sample_rich.docx"))


def _build_pdf(pages: Sequence[str]) -> bytes:
    """한국어 텍스트가 추출되는 최소 PDF를 손으로 조립한다.

    fpdf2 등으로 한국어를 그리려면 CJK TTF를 임베드해야 하는데, 배포 가능한 한글
    폰트가 저장소에 없고 시스템 폰트를 커밋된 픽스처에 심는 것은 라이선스 문제가 있다.
    그래서 폰트 프로그램 없이 ToUnicode CMap(PDF 1.7 §9.10.3)만 붙인다 — 뷰어에서
    글자 모양은 깨지지만 텍스트 추출 경로(pypdf가 ToUnicode로 문자 코드를 유니코드로
    되돌리는 경로)는 실제 한글 PDF와 같다.

    문자 코드는 1바이트(0x01~)로 배정하므로 픽스처 전체 고유 문자는 255자 이하여야 한다.
    """
    unique_chars = sorted({char for page in pages for char in page})
    if len(unique_chars) > 255:
        raise ValueError("1바이트 코드로 담을 수 있는 고유 문자 수를 넘었다")
    code_of = {char: index for index, char in enumerate(unique_chars, start=1)}

    objects: list[bytes] = []

    def add(body: bytes) -> int:
        """객체를 등록하고 1-기반 객체 번호를 돌려준다."""
        objects.append(body)
        return len(objects)

    def add_stream(payload: bytes) -> int:
        return add(b"<< /Length %d >>\nstream\n" % len(payload) + payload + b"\nendstream")

    # 객체 번호를 미리 잡아둔다 (Catalog가 Pages를, Pages가 Page들을 참조).
    catalog_number, pages_number = 1, 2
    objects.extend((b"", b""))

    bfchar_lines = "\n".join(f"<{code:02X}> <{ord(char):04X}>" for char, code in code_of.items())
    cmap = f"""/CIDInit /ProcSet findresource begin
12 dict begin
begincmap
/CMapName /EasyDocFixture def
/CMapType 2 def
1 begincodespacerange
<00> <FF>
endcodespacerange
{len(code_of)} beginbfchar
{bfchar_lines}
endbfchar
endcmap
CMapName currentdict /CMap defineresource pop
end
end""".encode()
    tounicode_number = add_stream(cmap) if code_of else 0
    font_body = b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica"
    if tounicode_number:
        font_body += b" /ToUnicode %d 0 R" % tounicode_number
    font_number = add(font_body + b" >>")

    page_numbers: list[int] = []
    for page in pages:
        codes = "".join(f"{code_of[char]:02X}" for char in page)
        content = f"BT /F1 12 Tf 72 720 Td <{codes}> Tj ET".encode() if page else b""
        content_number = add_stream(content)
        page_numbers.append(
            add(
                b"<< /Type /Page /Parent %d 0 R /MediaBox [0 0 595 842] "
                b"/Resources << /Font << /F1 %d 0 R >> >> /Contents %d 0 R >>"
                % (pages_number, font_number, content_number)
            )
        )

    objects[catalog_number - 1] = b"<< /Type /Catalog /Pages %d 0 R >>" % pages_number
    kids = b" ".join(b"%d 0 R" % number for number in page_numbers)
    objects[pages_number - 1] = b"<< /Type /Pages /Kids [%s] /Count %d >>" % (
        kids,
        len(page_numbers),
    )

    out = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for number, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += b"%d 0 obj\n" % number + body + b"\nendobj\n"
    xref_offset = len(out)
    out += b"xref\n0 %d\n0000000000 65535 f \n" % (len(objects) + 1)
    for offset in offsets:
        out += b"%010d 00000 n \n" % offset
    out += b"trailer\n<< /Size %d /Root %d 0 R >>\nstartxref\n%d\n%%%%EOF\n" % (
        len(objects) + 1,
        catalog_number,
        xref_offset,
    )
    return bytes(out)


def _write_pdfs() -> None:
    (FIXTURES_DIR / "sample.pdf").write_bytes(_build_pdf(PDF_PAGES))
    # 텍스트 레이어가 없는 PDF — 스캔 문서를 대신한다.
    (FIXTURES_DIR / "empty.pdf").write_bytes(_build_pdf(("",)))


def _hwpx_section(paragraphs: Sequence[str], *, split_first: bool) -> bytes:
    """OWPML 구역 XML을 만든다.

    구조 출처: KS X 6101 (OWPML) — 구역 루트 `<hs:sec>`, 문단 `<hp:p>`,
    글자 조각 `<hp:run>/<hp:t>`. 네임스페이스 URI는 한글 문서 파일 형식 공개 문서의
    hwpml/2011 네임스페이스를 따른다.
    """
    body: list[str] = []
    for index, text in enumerate(paragraphs):
        if split_first and index == 0:
            # 한 문단이 여러 run으로 쪼개진 경우 — 구분자 없이 이어 붙어야 한다.
            head, _, tail = text.partition(" ")
            runs = f"<hp:run><hp:t>{head} </hp:t></hp:run><hp:run><hp:t>{tail}</hp:t></hp:run>"
        else:
            runs = f"<hp:run><hp:t>{text}</hp:t></hp:run>"
        body.append(f"<hp:p>{runs}</hp:p>")
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section"'
        ' xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">'
        f"{''.join(body)}</hs:sec>"
    ).encode()


def _write_hwpx() -> None:
    """hwpx(= OWPML zip 패키지) 픽스처.

    한컴 오피스가 없어 실제 저장 파일을 쓸 수 없으므로 스펙대로 조립한다. 추출에
    필요한 `Contents/section*.xml` 외에 mimetype·META-INF·content.hpf를 함께 넣어
    실제 패키지 모양을 유지한다. **section1을 section0보다 먼저 기록**해, 추출기가
    zip 기록 순서가 아니라 구역 번호로 정렬하는지 확인한다.
    """
    path = FIXTURES_DIR / "sample.hwpx"
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        # mimetype은 스펙상 첫 항목이고 무압축이다 (ODF 패키지 규약과 같다).
        archive.writestr(zipfile.ZipInfo("mimetype"), b"application/hwp+zip", zipfile.ZIP_STORED)
        archive.writestr(
            "META-INF/container.xml",
            '<?xml version="1.0" encoding="UTF-8"?>'
            '<container xmlns="urn:oasis:names:tc:opendocument:xmlns:container">'
            '<rootfiles><rootfile full-path="Contents/content.hpf"'
            ' media-type="application/hwpml-package+xml"/></rootfiles></container>',
        )
        archive.writestr(
            "version.xml",
            '<?xml version="1.0" encoding="UTF-8"?>'
            '<hv:HCFVersion xmlns:hv="http://www.hancom.co.kr/hwpml/2011/version"'
            ' tagetApplication="WORDPROCESSOR" major="5" minor="1" micro="0" buildNumber="0"/>',
        )
        archive.writestr(
            "Contents/content.hpf",
            '<?xml version="1.0" encoding="UTF-8"?>'
            '<opf:package xmlns:opf="http://www.idpf.org/2007/opf/" version="">'
            '<opf:spine><opf:itemref idref="header"/>'
            '<opf:itemref idref="section0"/><opf:itemref idref="section1"/>'
            "</opf:spine></opf:package>",
        )
        archive.writestr("Contents/section1.xml", _hwpx_section(HWPX_SECTION1, split_first=False))
        archive.writestr("Contents/section0.xml", _hwpx_section(HWPX_SECTION0, split_first=True))


def main() -> None:
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)
    _write_docx_paragraphs()
    _write_docx_table()
    _write_docx_rich()
    _write_pdfs()
    _write_hwpx()
    for path in sorted(FIXTURES_DIR.iterdir()):
        print(f"생성: {path.name} ({path.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
