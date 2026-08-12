"""spike 전용 추가 픽스처 — 저장소 fixtures는 건드리지 않는다."""
import io, json, sys, pathlib, zipfile
sys.path.insert(0, "/Users/harris/Development/private/easy-doc")
import docx
from docx.oxml import parse_xml
from app.ingest.extractors import extract_text, _docx_blocks

SP = pathlib.Path("/private/tmp/claude-503/-Users-harris-Development-private-easy-doc/6f3e0698-996a-4c4a-b701-ef39bb65da0a/scratchpad/extra")
SP.mkdir(exist_ok=True)

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
A = "http://schemas.openxmlformats.org/drawingml/2006/main"
M = "http://schemas.openxmlformats.org/officeDocument/2006/math"
WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"

d = docx.Document()
d.add_paragraph("본문 첫 문단")

# 1) SDT (구조화 문서 태그) — 본문 수준 콘텐츠 컨트롤
d.element.body.insert(1, parse_xml(
    f'<w:sdt xmlns:w="{W}"><w:sdtPr><w:alias w:val="테스트컨트롤"/></w:sdtPr>'
    f'<w:sdtContent><w:p><w:r><w:t>SDT 안의 문장입니다.</w:t></w:r></w:p></w:sdtContent></w:sdt>'
))

# 2) a:t — DrawingML 도형 안 텍스트 (w:txbxContent 없이 a:t 직접)
d.add_paragraph()._p.append(parse_xml(
    f'<w:r xmlns:w="{W}" xmlns:wp="{WP}" xmlns:a="{A}">'
    f'<w:drawing><wp:inline><a:graphic><a:graphicData>'
    f'<a:p><a:r><a:t>도형 텍스트입니다.</a:t></a:r></a:p>'
    f'</a:graphicData></a:graphic></wp:inline></w:drawing></w:r>'
))

# 3) m:t — OMML 수식
d.add_paragraph()._p.append(parse_xml(
    f'<m:oMath xmlns:m="{M}"><m:r><m:t>x+1=2</m:t></m:r></m:oMath>'
))

d.add_paragraph("마지막 문단")
p = SP / "sdt_shape_math.docx"
d.save(str(p))

data = p.read_bytes()
out = {
    "sdt_shape_math.docx": {"text": extract_text("a.docx", data)},
    "sdt_shape_math.docx::blocks": list(_docx_blocks(docx.Document(io.BytesIO(data)))),
}

# 4) 비정상 PDF — 다단·역순 배치. pypdf vs PDFBox 차이를 드러내기 위한 합성 문서.
def build_layout_pdf():
    txt = [
        (300, 700, "오른쪽 단 첫 줄"),
        (72,  700, "왼쪽 단 첫 줄"),
        (72,  680, "왼쪽 단 둘째 줄"),
        (300, 680, "오른쪽 단 둘째 줄"),
    ]
    chars = sorted({c for _,_,s in txt for c in s})
    code = {c:i for i,c in enumerate(chars, start=1)}
    objs = []
    def add(b):
        objs.append(b); return len(objs)
    def add_stream(p):
        return add(b"<< /Length %d >>\nstream\n" % len(p) + p + b"\nendstream")
    cat, pages_n = 1, 2
    objs.extend((b"", b""))
    bf = "\n".join(f"<{v:02X}> <{ord(k):04X}>" for k,v in code.items())
    cmap = f"""/CIDInit /ProcSet findresource begin
12 dict begin
begincmap
/CMapName /L def
/CMapType 2 def
1 begincodespacerange
<00> <FF>
endcodespacerange
{len(code)} beginbfchar
{bf}
endbfchar
endcmap
CMapName currentdict /CMap defineresource pop
end
end""".encode()
    tu = add_stream(cmap)
    fn = add(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica /ToUnicode %d 0 R >>" % tu)
    parts = []
    for x, y, s in txt:
        codes = "".join(f"{code[c]:02X}" for c in s)
        parts.append(f"BT /F1 12 Tf {x} {y} Td <{codes}> Tj ET")
    cn = add_stream(" ".join(parts).encode())
    pg = add(b"<< /Type /Page /Parent %d 0 R /MediaBox [0 0 595 842] /Resources << /Font << /F1 %d 0 R >> >> /Contents %d 0 R >>" % (pages_n, fn, cn))
    objs[cat-1] = b"<< /Type /Catalog /Pages %d 0 R >>" % pages_n
    objs[pages_n-1] = b"<< /Type /Pages /Kids [%d 0 R] /Count 1 >>" % pg
    out = bytearray(b"%PDF-1.4\n"); offs=[]
    for n, b in enumerate(objs, start=1):
        offs.append(len(out)); out += b"%d 0 obj\n" % n + b + b"\nendobj\n"
    xo = len(out)
    out += b"xref\n0 %d\n0000000000 65535 f \n" % (len(objs)+1)
    for o in offs: out += b"%010d 00000 n \n" % o
    out += b"trailer\n<< /Size %d /Root %d 0 R >>\nstartxref\n%d\n%%%%EOF\n" % (len(objs)+1, cat, xo)
    return bytes(out)

lp = SP / "layout.pdf"
lp.write_bytes(build_layout_pdf())
out["layout.pdf"] = {"text": extract_text("l.pdf", lp.read_bytes())}

# 5) 선언 크기를 위조한 zip (압축 폭탄 + 헤더 거짓말)
payload = b"\0" * (80 * 1024 * 1024)
buf = io.BytesIO()
with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
    z.writestr("Contents/section0.xml", payload)
raw = bytearray(buf.getvalue())
# 로컬 헤더(offset 22)와 중앙 디렉터리의 uncompressed size 필드를 1024로 위조
import struct
raw[22:26] = struct.pack("<I", 1024)
cd = raw.rfind(b"PK\x01\x02")
raw[cd+24:cd+28] = struct.pack("<I", 1024)
(SP / "forged_size.zip").write_bytes(bytes(raw))
try:
    extract_text("f.hwpx", bytes(raw))
    out["forged_size"] = {"python": "통과(!!)"}
except Exception as e:
    out["forged_size"] = {"python": f"{type(e).__name__}: {e}"}

(pathlib.Path("/private/tmp/claude-503/-Users-harris-Development-private-easy-doc/6f3e0698-996a-4c4a-b701-ef39bb65da0a/scratchpad/extra_oracle.json")).write_text(json.dumps(out, ensure_ascii=False, indent=2))
print(json.dumps(out, ensure_ascii=False, indent=2))
