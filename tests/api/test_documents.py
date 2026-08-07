"""문서 API 테스트.

DB·Redis 없이 돈다 — 저장소·큐·암호기 의존성만 fake로 갈아끼우고 라우터·서비스·
추출기·예외 핸들러는 실제 코드를 그대로 태운다.

여기서 지키려는 불변식은 네 가지다: 저장 전 암호화, INSERT→commit→enqueue 순서,
소유자 격리, 그리고 응답·오류에 본문이 새지 않는 것.
"""

import io
import uuid
from collections.abc import Iterator
from contextlib import contextmanager
from datetime import UTC, datetime, timedelta
from pathlib import Path
from typing import Any
from urllib.parse import quote

import docx
import jwt
import pytest
from cryptography.fernet import Fernet
from fastapi.testclient import TestClient
from pydantic import SecretStr

from app.api.deps import (
    get_conversion_repository,
    get_document_repository,
    get_settings,
    get_task_queue,
    get_text_cipher,
    get_user_repository,
)
from app.config import Settings
from app.ingest.extractors import MAX_UPLOAD_BYTES
from app.main import app
from app.models.conversion import ConversionStatus
from app.models.user import User
from app.privacy.crypto import TextCipher
from app.privacy.masking import MaskCategory, MaskedItem
from app.services.documents import MAX_CONVERTIBLE_CHARS, serialize_masked_items
from tests.fakes import FakeConversionStore, FakeDocumentStore, FakeTaskQueue, FakeUserRepository

_SECRET = "test-secret-do-not-use-in-production"
_FIXTURES = Path(__file__).resolve().parent.parent / "ingest" / "fixtures"
_TEXT = "재난지원금 안내\n신청 기간은 3월 2일부터입니다."
#: 워커가 남긴 AI 초안(자리표시자 없는 기본형).
_EASY = "신청 기간은 3월 2일부터예요."


class Fixture:
    """테스트 한 건이 쓰는 대역 묶음 (저장소·큐·암호기·사용자)."""

    def __init__(self, *, queue_error: Exception | None = None) -> None:
        self.journal: list[str] = []
        self.documents = FakeDocumentStore(self.journal)
        self.conversions = FakeConversionStore(self.documents)
        self.queue = FakeTaskQueue(self.journal, error=queue_error)
        self.cipher = TextCipher(Fernet.generate_key().decode())
        self.users = FakeUserRepository()
        self.user = self._register("owner@example.com")

    def _register(self, email: str) -> User:
        """가입 흐름(argon2)을 거치지 않고 사용자를 넣는다."""
        user = User(
            id=uuid.uuid4(),
            email=email,
            password_hash="$argon2id$fake",
            created_at=datetime.now(UTC),
        )
        self.users.add(user)
        return user

    def other_user(self) -> User:
        """다른 소유자 — 격리 테스트용."""
        return self._register("stranger@example.com")

    def headers(self, user: User | None = None) -> dict[str, str]:
        """Bearer 인증 헤더. 토큰은 AuthService가 발급하는 것과 같은 모양이다."""
        claims = {
            "sub": str((user or self.user).id),
            "exp": datetime.now(UTC) + timedelta(minutes=10),
            "typ": "access",
        }
        return {"Authorization": f"Bearer {jwt.encode(claims, _SECRET, 'HS256')}"}


def _settings(fernet_key: str | None = None) -> Settings:
    """테스트용 설정. 생성자 인자가 .env·환경변수보다 우선한다.

    fernet_key는 "설정됨/안 됨"만 보면 되므로 실제 키 형식이 아니어도 된다 —
    암호기는 대역으로 갈아끼우고, 키 값 자체를 쓰는 경로는 여기 없다.
    """
    return Settings(
        jwt_secret=SecretStr(_SECRET),
        fernet_key=SecretStr(fernet_key) if fernet_key is not None else None,
    )


@contextmanager
def _client_with(fixture: Fixture, *, queue_ready: bool = True) -> Iterator[TestClient]:
    """대역을 주입한 테스트 클라이언트.

    `queue_ready=False`는 큐 의존성을 갈아끼우는 대신 lifespan이 남긴 상태를 비운다 —
    "Redis에 붙지 못한 채 뜬 앱"을 실제 `get_task_queue`가 어떻게 다루는지 봐야 한다.
    """
    app.dependency_overrides[get_user_repository] = lambda: fixture.users
    app.dependency_overrides[get_document_repository] = lambda: fixture.documents
    app.dependency_overrides[get_conversion_repository] = lambda: fixture.conversions
    app.dependency_overrides[get_text_cipher] = lambda: fixture.cipher
    app.dependency_overrides[get_settings] = lambda: _settings("x" * 44)
    if queue_ready:
        app.dependency_overrides[get_task_queue] = lambda: fixture.queue
    try:
        with TestClient(app) as test_client:
            if not queue_ready:
                app.state.task_queue = None
            yield test_client
    finally:
        app.dependency_overrides.clear()


@pytest.fixture
def fixture() -> Fixture:
    """테스트마다 비어 있는 대역 묶음."""
    return Fixture()


@pytest.fixture
def client(fixture: Fixture) -> Iterator[TestClient]:
    """정상 설정으로 구성한 클라이언트."""
    with _client_with(fixture) as test_client:
        yield test_client


def _upload_text(
    client: TestClient, fixture: Fixture, text: str = _TEXT, **extra: Any
) -> dict[str, Any]:
    """텍스트 모드로 업로드하고 응답 본문을 돌려준다(성공 전제)."""
    response = client.post("/documents", json={"text": text, **extra}, headers=fixture.headers())
    assert response.status_code == 202, response.text
    body: dict[str, Any] = response.json()
    return body


def _empty_docx() -> bytes:
    """문단이 하나도 없는 docx — 추출은 성공하지만 결과가 빈 문자열이다."""
    buffer = io.BytesIO()
    docx.Document().save(buffer)
    return buffer.getvalue()


# --- 업로드 정상 흐름 -----------------------------------------------------------


def test_텍스트를_올리면_변환이_접수된다(client: TestClient, fixture: Fixture) -> None:
    body = _upload_text(client, fixture)

    assert body["status"] == ConversionStatus.PENDING
    assert body["char_count"] == len(_TEXT)
    assert uuid.UUID(body["document_id"])
    # 큐에는 변환 식별자만 넘긴다 — 본문이 Redis에 남지 않는다.
    assert fixture.queue.jobs == [("convert_document", (body["conversion_id"],))]
    # 작업 id를 변환 id로 고정해 중복 등록을 arq가 걸러내게 한다.
    assert fixture.queue.job_ids == [body["conversion_id"]]


def test_접수_응답이_결과_주소를_알려준다(client: TestClient, fixture: Fixture) -> None:
    """202는 "나중에 여기서 확인하라"는 뜻이다 — 그 자리를 표준 헤더로 알려 준다."""
    response = client.post("/documents", json={"text": _TEXT}, headers=fixture.headers())

    assert response.headers["Location"] == f"/conversions/{response.json()['conversion_id']}"


def test_저장을_확정한_뒤에_큐에_넣는다(client: TestClient, fixture: Fixture) -> None:
    """커밋 전에 넣으면 워커가 아직 보이지 않는 행을 읽으러 가서 곧바로 실패한다."""
    _upload_text(client, fixture)

    assert fixture.journal == ["create_document", "create_conversion", "commit", "enqueue"]


def test_원문은_암호화해서_저장한다(client: TestClient, fixture: Fixture) -> None:
    body = _upload_text(client, fixture)

    document = fixture.documents.documents[uuid.UUID(body["document_id"])]
    assert "재난지원금".encode() not in document.source_text_encrypted
    assert fixture.cipher.decrypt(document.source_text_encrypted) == _TEXT


def test_제목이_없으면_본문_첫_줄에서_유도한다(client: TestClient, fixture: Fixture) -> None:
    body = _upload_text(client, fixture)

    document = fixture.documents.documents[uuid.UUID(body["document_id"])]
    assert document.title == "재난지원금 안내"


def test_제목을_주면_그대로_쓴다(client: TestClient, fixture: Fixture) -> None:
    body = _upload_text(client, fixture, title="  3월 안내문  ")

    document = fixture.documents.documents[uuid.UUID(body["document_id"])]
    assert document.title == "3월 안내문"


def test_첫_줄이_없으면_기본_제목을_쓴다(client: TestClient, fixture: Fixture) -> None:
    body = _upload_text(client, fixture, text="   \n\n   내용")

    document = fixture.documents.documents[uuid.UUID(body["document_id"])]
    assert document.title == "내용"


def test_파일을_올리면_본문을_뽑아_저장한다(client: TestClient, fixture: Fixture) -> None:
    data = (_FIXTURES / "sample.docx").read_bytes()

    response = client.post(
        "/documents",
        files={"file": ("안내문.docx", data, "application/octet-stream")},
        headers=fixture.headers(),
    )

    assert response.status_code == 202, response.text
    document = fixture.documents.documents[uuid.UUID(response.json()["document_id"])]
    assert document.source_format == "docx"
    assert document.char_count > 0
    # 제목은 파일명이 아니라 본문에서 유도한다 — 파일명 자체가 개인정보일 수 있다
    # (`홍길동_주민등록등본.pdf`). 파일명은 어디에도 저장하지 않는다.
    assert ".docx" not in document.title
    assert fixture.cipher.decrypt(document.source_text_encrypted).startswith(document.title)


def test_파일_모드에서도_제목을_받는다(client: TestClient, fixture: Fixture) -> None:
    data = (_FIXTURES / "sample.docx").read_bytes()

    response = client.post(
        "/documents",
        files={"file": ("a.docx", data, "application/octet-stream")},
        data={"title": "복지 안내"},
        headers=fixture.headers(),
    )

    assert response.status_code == 202, response.text
    document = fixture.documents.documents[uuid.UUID(response.json()["document_id"])]
    assert document.title == "복지 안내"


# --- 업로드 경계 ---------------------------------------------------------------


@pytest.mark.parametrize("text", ["", "   \n\t  "])
def test_빈_본문은_422(client: TestClient, fixture: Fixture, text: str) -> None:
    response = client.post("/documents", json={"text": text}, headers=fixture.headers())

    assert response.status_code == 422


def test_본문_필드가_없으면_422(client: TestClient, fixture: Fixture) -> None:
    response = client.post("/documents", json={"title": "제목만"}, headers=fixture.headers())

    assert response.status_code == 422
    assert "text" in response.text


def test_JSON이_아니면_422(client: TestClient, fixture: Fixture) -> None:
    response = client.post(
        "/documents",
        content=b"not json",
        headers={**fixture.headers(), "Content-Type": "application/json"},
    )

    assert response.status_code == 422


def test_변환_상한을_넘는_본문은_422(client: TestClient, fixture: Fixture) -> None:
    """길이는 바이트가 아니라 문자 수로 잰다 (한국어는 글자당 3바이트다).

    상한을 넘는 문서는 LLM 비용을 다 치른 뒤 절단으로 실패하는 것이 확정이므로,
    돈을 쓰기 전에 거절한다.
    """
    response = client.post(
        "/documents", json={"text": "가" * (MAX_CONVERTIBLE_CHARS + 1)}, headers=fixture.headers()
    )

    assert response.status_code == 422
    assert "분할 변환" in response.json()["detail"]
    # 저장도 큐 등록도 일어나지 않는다.
    assert fixture.journal == []


def test_상한_길이_문서는_통과한다(client: TestClient, fixture: Fixture) -> None:
    """경계에서 한 글자 차이로 정상 문서를 거절하지 않는지 본다."""
    response = client.post(
        "/documents", json={"text": "가" * MAX_CONVERTIBLE_CHARS}, headers=fixture.headers()
    )

    assert response.status_code == 202, response.text


def test_변환_상한은_파일_업로드에도_적용된다(client: TestClient, fixture: Fixture) -> None:
    """입력 방식에 따라 변환 가능 여부가 달라지면 사용자에게 설명할 수 없다."""
    document = docx.Document()
    document.add_paragraph("가" * (MAX_CONVERTIBLE_CHARS + 1))
    buffer = io.BytesIO()
    document.save(buffer)

    response = client.post(
        "/documents",
        files={"file": ("long.docx", buffer.getvalue(), "application/octet-stream")},
        headers=fixture.headers(),
    )

    assert response.status_code == 422
    assert "분할 변환" in response.json()["detail"]
    assert fixture.journal == []


def test_상한을_넘는_파일은_413(client: TestClient, fixture: Fixture) -> None:
    oversized = b"a" * (MAX_UPLOAD_BYTES + 1)

    response = client.post(
        "/documents",
        files={"file": ("big.docx", oversized, "application/octet-stream")},
        headers=fixture.headers(),
    )

    assert response.status_code == 413
    assert fixture.journal == []


def test_지원하지_않는_형식은_422(client: TestClient, fixture: Fixture) -> None:
    response = client.post(
        "/documents",
        files={"file": ("old.hwp", b"binary", "application/octet-stream")},
        headers=fixture.headers(),
    )

    assert response.status_code == 422
    assert "docx" in response.json()["detail"]


def test_추출_결과가_비면_422(client: TestClient, fixture: Fixture) -> None:
    """빈 docx는 예외 없이 ""를 돌려준다 — 빈 문서 판정은 추출 결과로 해야 한다."""
    response = client.post(
        "/documents",
        files={"file": ("empty.docx", _empty_docx(), "application/octet-stream")},
        headers=fixture.headers(),
    )

    assert response.status_code == 422
    assert fixture.journal == []


def test_텍스트_레이어가_없는_PDF는_422(client: TestClient, fixture: Fixture) -> None:
    data = (_FIXTURES / "empty.pdf").read_bytes()

    response = client.post(
        "/documents",
        files={"file": ("scan.pdf", data, "application/octet-stream")},
        headers=fixture.headers(),
    )

    assert response.status_code == 422
    assert "스캔" in response.json()["detail"]


def test_파일_파트가_없는_multipart는_422(client: TestClient, fixture: Fixture) -> None:
    response = client.post("/documents", data={"title": "제목"}, headers=fixture.headers())

    assert response.status_code == 422


# --- 큐·설정 장애 ---------------------------------------------------------------


def test_큐_등록에_실패하면_502이고_변환은_실패로_남는다(fixture: Fixture) -> None:
    """커밋은 이미 끝났다 — 그대로 두면 변환이 영원히 pending으로 굳는다."""
    fixture = Fixture(queue_error=ConnectionError("redis down"))

    with _client_with(fixture) as client:
        response = client.post("/documents", json={"text": _TEXT}, headers=fixture.headers())

    assert response.status_code == 502
    # 예외 메시지(접속 정보가 실릴 수 있다)가 응답으로 새면 안 된다.
    assert "redis down" not in response.text
    conversion = next(iter(fixture.conversions.conversions.values()))
    assert conversion.status == ConversionStatus.FAILED
    assert conversion.failure_code == "EnqueueFailed"


def test_큐가_준비되지_않았으면_503(fixture: Fixture) -> None:
    with _client_with(fixture, queue_ready=False) as client:
        response = client.post("/documents", json={"text": _TEXT}, headers=fixture.headers())

    assert response.status_code == 503


def test_암호화_키가_없으면_503(fixture: Fixture) -> None:
    """운영 설정 누락은 사용자 잘못이 아니다 — 4xx가 아니라 5xx로 알린다."""
    app.dependency_overrides[get_user_repository] = lambda: fixture.users
    app.dependency_overrides[get_settings] = lambda: _settings(fernet_key=None)
    app.dependency_overrides[get_task_queue] = lambda: fixture.queue
    try:
        with TestClient(app) as client:
            response = client.post("/documents", json={"text": _TEXT}, headers=fixture.headers())
    finally:
        app.dependency_overrides.clear()

    assert response.status_code == 503


# --- 변환 조회 -----------------------------------------------------------------


def _complete(fixture: Fixture, conversion_id: str, *, easy_text: str = _EASY) -> None:
    """워커가 성공적으로 끝낸 상태를 만든다 (암호화 저장 형태 그대로)."""
    conversion = fixture.conversions.conversions[uuid.UUID(conversion_id)]
    items = [
        MaskedItem(
            category=MaskCategory.PHONE,
            placeholder="[[전화번호1]]",
            original=SecretStr("010-1234-5678"),
        )
    ]
    conversion.status = ConversionStatus.DONE
    conversion.easy_text_encrypted = fixture.cipher.encrypt(easy_text)
    conversion.masked_items_encrypted = fixture.cipher.encrypt(serialize_masked_items(items))
    conversion.missing_placeholders = ["[[이메일1]]"]
    conversion.provider_name = "fake"
    conversion.model = "fake-model"
    conversion.input_tokens = 10
    conversion.output_tokens = 20


def test_대기_중이면_상태만_돌려준다(client: TestClient, fixture: Fixture) -> None:
    conversion_id = _upload_text(client, fixture)["conversion_id"]

    body = client.get(f"/conversions/{conversion_id}", headers=fixture.headers()).json()

    assert body["status"] == ConversionStatus.PENDING
    assert body["easy_text"] is None
    assert body["masked_items"] == []
    assert body["failure_code"] is None


def test_완료되면_결과와_마스킹_항목을_복호화해_돌려준다(
    client: TestClient, fixture: Fixture
) -> None:
    conversion_id = _upload_text(client, fixture)["conversion_id"]
    _complete(fixture, conversion_id)

    body = client.get(f"/conversions/{conversion_id}", headers=fixture.headers()).json()

    assert body["status"] == ConversionStatus.DONE
    assert body["easy_text"] == "신청 기간은 3월 2일부터예요."
    # 검수 화면은 원문 대조가 목적이라 원문을 함께 본다 (소유자 인증 필수).
    assert body["masked_items"] == [
        {
            "category": "전화번호",
            "placeholder": "[[전화번호1]]",
            "original": "010-1234-5678",
        }
    ]
    assert body["missing_placeholders"] == ["[[이메일1]]"]
    assert (body["provider_name"], body["model"]) == ("fake", "fake-model")
    assert (body["input_tokens"], body["output_tokens"]) == (10, 20)


def test_실패하면_사유_코드를_돌려준다(client: TestClient, fixture: Fixture) -> None:
    conversion_id = _upload_text(client, fixture)["conversion_id"]
    conversion = fixture.conversions.conversions[uuid.UUID(conversion_id)]
    conversion.status = ConversionStatus.FAILED
    conversion.failure_code = "LLMTruncatedError"

    body = client.get(f"/conversions/{conversion_id}", headers=fixture.headers()).json()

    assert body["status"] == ConversionStatus.FAILED
    assert body["failure_code"] == "LLMTruncatedError"
    assert body["easy_text"] is None


def test_남의_변환은_404(client: TestClient, fixture: Fixture) -> None:
    """있다는 사실 자체를 알리지 않는다 — 403이 아니라 404다."""
    conversion_id = _upload_text(client, fixture)["conversion_id"]
    stranger = fixture.other_user()

    response = client.get(f"/conversions/{conversion_id}", headers=fixture.headers(stranger))

    assert response.status_code == 404


def test_없는_변환은_404(client: TestClient, fixture: Fixture) -> None:
    response = client.get(f"/conversions/{uuid.uuid4()}", headers=fixture.headers())

    assert response.status_code == 404


# --- 검수 수정 저장 -------------------------------------------------------------

_EDITED = "신청 기간은 3월 2일부터 3월 31일까지예요."


def _review(
    client: TestClient,
    fixture: Fixture,
    conversion_id: str,
    *,
    text: str = _EDITED,
    user: User | None = None,
) -> Any:
    """검수 수정본을 저장한다(응답을 그대로 돌려준다 — 실패 경로도 본다)."""
    return client.put(
        f"/conversions/{conversion_id}",
        json={"edited_text": text},
        headers=fixture.headers(user),
    )


def test_검수_수정본을_저장하면_조회에_함께_나온다(client: TestClient, fixture: Fixture) -> None:
    conversion_id = _upload_text(client, fixture)["conversion_id"]
    _complete(fixture, conversion_id)

    saved = _review(client, fixture, conversion_id)

    assert saved.status_code == 200, saved.text
    assert saved.json()["edited_text"] == _EDITED
    assert saved.json()["reviewed_at"] is not None
    body = client.get(f"/conversions/{conversion_id}", headers=fixture.headers()).json()
    assert body["edited_text"] == _EDITED
    assert body["reviewed_at"] == saved.json()["reviewed_at"]


def test_검수_전에는_수정본_필드가_비어_있다(client: TestClient, fixture: Fixture) -> None:
    conversion_id = _upload_text(client, fixture)["conversion_id"]
    _complete(fixture, conversion_id)

    body = client.get(f"/conversions/{conversion_id}", headers=fixture.headers()).json()

    assert body["edited_text"] is None
    assert body["reviewed_at"] is None


def test_검수_수정본은_AI_초안을_덮지_않는다(client: TestClient, fixture: Fixture) -> None:
    """초안은 수정률 KPI(master-plan 7장)의 기준선이다 — 덮어쓰면 되살릴 수 없다."""
    conversion_id = _upload_text(client, fixture)["conversion_id"]
    _complete(fixture, conversion_id)

    body = _review(client, fixture, conversion_id).json()

    assert body["easy_text"] == "신청 기간은 3월 2일부터예요."
    conversion = fixture.conversions.conversions[uuid.UUID(conversion_id)]
    assert conversion.easy_text_encrypted is not None
    assert fixture.cipher.decrypt(conversion.easy_text_encrypted) == "신청 기간은 3월 2일부터예요."


def test_검수_수정본은_암호화해서_저장한다(client: TestClient, fixture: Fixture) -> None:
    conversion_id = _upload_text(client, fixture)["conversion_id"]
    _complete(fixture, conversion_id)

    _review(client, fixture, conversion_id)

    conversion = fixture.conversions.conversions[uuid.UUID(conversion_id)]
    assert conversion.edited_text_encrypted is not None
    assert "신청 기간".encode() not in conversion.edited_text_encrypted
    assert fixture.cipher.decrypt(conversion.edited_text_encrypted) == _EDITED


def test_다시_수정하면_덮어쓴다(client: TestClient, fixture: Fixture) -> None:
    conversion_id = _upload_text(client, fixture)["conversion_id"]
    _complete(fixture, conversion_id)
    _review(client, fixture, conversion_id)

    body = _review(client, fixture, conversion_id, text="3월 2일부터 신청하세요.").json()

    assert body["edited_text"] == "3월 2일부터 신청하세요."


def test_완료되지_않은_변환은_수정할_수_없다(client: TestClient, fixture: Fixture) -> None:
    """아직 결과가 없는데 수정본을 받으면 무엇을 고친 것인지 설명할 수 없다."""
    conversion_id = _upload_text(client, fixture)["conversion_id"]

    response = _review(client, fixture, conversion_id)

    assert response.status_code == 409
    assert fixture.conversions.conversions[uuid.UUID(conversion_id)].edited_text_encrypted is None


def test_실패한_변환도_수정할_수_없다(client: TestClient, fixture: Fixture) -> None:
    conversion_id = _upload_text(client, fixture)["conversion_id"]
    conversion = fixture.conversions.conversions[uuid.UUID(conversion_id)]
    conversion.status = ConversionStatus.FAILED
    conversion.failure_code = "LLMProviderError"

    assert _review(client, fixture, conversion_id).status_code == 409


def test_남의_변환은_수정할_수_없다(client: TestClient, fixture: Fixture) -> None:
    """소유자 격리 — 있다는 사실 자체를 알리지 않는다(403이 아니라 404)."""
    conversion_id = _upload_text(client, fixture)["conversion_id"]
    _complete(fixture, conversion_id)
    stranger = fixture.other_user()

    response = _review(client, fixture, conversion_id, user=stranger)

    assert response.status_code == 404
    assert fixture.conversions.conversions[uuid.UUID(conversion_id)].edited_text_encrypted is None


def test_없는_변환_수정은_404(client: TestClient, fixture: Fixture) -> None:
    assert _review(client, fixture, str(uuid.uuid4())).status_code == 404


@pytest.mark.parametrize("text", ["", "   \n\t  "])
def test_빈_수정본은_422(client: TestClient, fixture: Fixture, text: str) -> None:
    conversion_id = _upload_text(client, fixture)["conversion_id"]
    _complete(fixture, conversion_id)

    assert _review(client, fixture, conversion_id, text=text).status_code == 422


def test_상한을_넘는_수정본은_422(client: TestClient, fixture: Fixture) -> None:
    """검수는 초안을 다듬는 자리다 — 원문과 같은 길이 상한을 그대로 적용한다."""
    conversion_id = _upload_text(client, fixture)["conversion_id"]
    _complete(fixture, conversion_id)

    response = _review(client, fixture, conversion_id, text="가" * (MAX_CONVERTIBLE_CHARS + 1))

    assert response.status_code == 422
    assert fixture.conversions.conversions[uuid.UUID(conversion_id)].edited_text_encrypted is None


def test_수정본_필드가_없으면_422(client: TestClient, fixture: Fixture) -> None:
    conversion_id = _upload_text(client, fixture)["conversion_id"]
    _complete(fixture, conversion_id)

    response = client.put(
        f"/conversions/{conversion_id}", json={"text": _EDITED}, headers=fixture.headers()
    )

    assert response.status_code == 422


# --- 내보내기 -----------------------------------------------------------------

#: 자리표시자가 남아 있는 초안 — 내보내기에서만 원문으로 복원된다.
_MASKED_EASY = "문의는 [[전화번호1]]로 하세요.\n\n신청은 3월 2일부터예요."


def _export(
    client: TestClient,
    fixture: Fixture,
    conversion_id: str,
    *,
    export_format: str = "docx",
    user: User | None = None,
) -> Any:
    """내보내기를 요청한다(응답을 그대로 돌려준다 — 실패 경로도 본다)."""
    return client.get(
        f"/conversions/{conversion_id}/export",
        params={"format": export_format},
        headers=fixture.headers(user),
    )


def _docx_text(content: bytes) -> str:
    """내려받은 docx를 다시 열어 전체 텍스트를 합친다."""
    return "\n".join(paragraph.text for paragraph in docx.Document(io.BytesIO(content)).paragraphs)


def test_docx로_내려받으면_자리표시자가_원문으로_복원된다(
    client: TestClient, fixture: Fixture
) -> None:
    """내보내기는 담당자가 그대로 배포할 최종 산출물이다 — 여기서만 마스킹을 되돌린다."""
    conversion_id = _upload_text(client, fixture)["conversion_id"]
    _complete(fixture, conversion_id, easy_text=_MASKED_EASY)

    response = _export(client, fixture, conversion_id)

    assert response.status_code == 200, response.text
    text = _docx_text(response.content)
    assert "[[" not in text
    assert "010-1234-5678" in text
    # 제목 + 빈 줄로 나뉜 문단 두 개.
    assert text == "재난지원금 안내\n문의는 010-1234-5678로 하세요.\n신청은 3월 2일부터예요."


def test_docx_응답_헤더가_형식과_파일명을_알려준다(client: TestClient, fixture: Fixture) -> None:
    conversion_id = _upload_text(client, fixture)["conversion_id"]
    _complete(fixture, conversion_id)

    response = _export(client, fixture, conversion_id)

    assert response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    disposition = response.headers["content-disposition"]
    # 한글 파일명은 RFC 5987(퍼센트 인코딩)으로만 실을 수 있다.
    assert f"filename*=UTF-8''{quote('재난지원금 안내.docx', safe='')}" in disposition


def test_txt로_내려받으면_UTF_8_본문이다(client: TestClient, fixture: Fixture) -> None:
    conversion_id = _upload_text(client, fixture)["conversion_id"]
    _complete(fixture, conversion_id, easy_text=_MASKED_EASY)

    response = _export(client, fixture, conversion_id, export_format="txt")

    assert response.status_code == 200, response.text
    assert response.headers["content-type"] == "text/plain; charset=utf-8"
    assert not response.content.startswith(b"\xef\xbb\xbf")
    assert response.content.decode("utf-8") == (
        "문의는 010-1234-5678로 하세요.\n\n신청은 3월 2일부터예요."
    )


def test_검수본이_있으면_검수본을_내보낸다(client: TestClient, fixture: Fixture) -> None:
    """담당자가 고친 뒤라면 그 결과가 최종본이다 (edited ?? easy)."""
    conversion_id = _upload_text(client, fixture)["conversion_id"]
    _complete(fixture, conversion_id, easy_text=_MASKED_EASY)
    _review(client, fixture, conversion_id, text="문의는 [[전화번호1]]로 주세요.")

    response = _export(client, fixture, conversion_id, export_format="txt")

    # 검수본에 남아 있던 자리표시자도 함께 복원된다.
    assert response.content.decode("utf-8") == "문의는 010-1234-5678로 주세요."


def test_완료되지_않은_변환은_내려받을_수_없다(client: TestClient, fixture: Fixture) -> None:
    conversion_id = _upload_text(client, fixture)["conversion_id"]

    assert _export(client, fixture, conversion_id).status_code == 409


def test_남의_변환은_내려받을_수_없다(client: TestClient, fixture: Fixture) -> None:
    conversion_id = _upload_text(client, fixture)["conversion_id"]
    _complete(fixture, conversion_id)
    stranger = fixture.other_user()

    response = _export(client, fixture, conversion_id, user=stranger)

    assert response.status_code == 404
    # 원문 개인정보가 남의 손에 넘어가는 경로다 — 본문 조각도 실리면 안 된다.
    assert "010-1234-5678" not in response.text


def test_없는_변환_내려받기는_404(client: TestClient, fixture: Fixture) -> None:
    assert _export(client, fixture, str(uuid.uuid4())).status_code == 404


@pytest.mark.parametrize("export_format", ["pdf", "hwp", ""])
def test_지원하지_않는_내보내기_형식은_422(
    client: TestClient, fixture: Fixture, export_format: str
) -> None:
    """pdf·hwp 내보내기는 Lean MVP 범위 밖이다 — 조용히 다른 형식을 주지 않는다."""
    conversion_id = _upload_text(client, fixture)["conversion_id"]
    _complete(fixture, conversion_id)

    assert _export(client, fixture, conversion_id, export_format=export_format).status_code == 422


def test_형식을_지정하지_않으면_422(client: TestClient, fixture: Fixture) -> None:
    conversion_id = _upload_text(client, fixture)["conversion_id"]
    _complete(fixture, conversion_id)

    response = client.get(f"/conversions/{conversion_id}/export", headers=fixture.headers())

    assert response.status_code == 422


# --- 목록 ---------------------------------------------------------------------


def test_목록은_문서_메타와_최신_변환_상태를_준다(client: TestClient, fixture: Fixture) -> None:
    first = _upload_text(client, fixture, text="첫 문서")
    second = _upload_text(client, fixture, text="둘째 문서")
    _complete(fixture, second["conversion_id"])

    body = client.get("/documents", headers=fixture.headers()).json()

    assert [item["id"] for item in body["items"]] == [
        second["document_id"],
        first["document_id"],
    ]
    assert body["items"][0]["status"] == ConversionStatus.DONE
    assert body["items"][0]["conversion_id"] == second["conversion_id"]
    assert body["items"][0]["source_format"] == "text"
    assert (body["limit"], body["offset"]) == (20, 0)
    assert body["has_more"] is False


def test_목록_응답에_본문이_실리지_않는다(client: TestClient, fixture: Fixture) -> None:
    _upload_text(client, fixture)

    response = client.get("/documents", headers=fixture.headers())

    assert "재난지원금 안내" in response.text  # 제목은 나온다
    assert "신청 기간은 3월 2일부터입니다" not in response.text
    assert "source_text_encrypted" not in response.text


def test_목록에_남의_문서는_섞이지_않는다(client: TestClient, fixture: Fixture) -> None:
    _upload_text(client, fixture)
    stranger = fixture.other_user()

    body = client.get("/documents", headers=fixture.headers(stranger)).json()

    assert body["items"] == []


def test_목록은_limit과_offset을_따른다(client: TestClient, fixture: Fixture) -> None:
    for index in range(3):
        _upload_text(client, fixture, text=f"문서 {index}")

    first = client.get("/documents?limit=2&offset=0", headers=fixture.headers()).json()
    body = client.get("/documents?limit=2&offset=2", headers=fixture.headers()).json()

    assert len(body["items"]) == 1
    assert (body["limit"], body["offset"]) == (2, 2)
    # 다음 페이지 유무는 한 건 더 읽어 판정한다 (COUNT 쿼리 없이).
    assert first["has_more"] is True
    assert body["has_more"] is False


@pytest.mark.parametrize("query", ["limit=0", "limit=101", "offset=-1"])
def test_범위를_벗어난_페이지_인자는_422(client: TestClient, fixture: Fixture, query: str) -> None:
    response = client.get(f"/documents?{query}", headers=fixture.headers())

    assert response.status_code == 422


# --- 인증 ---------------------------------------------------------------------


@pytest.mark.parametrize(
    ("method", "path"),
    [
        ("post", "/documents"),
        ("get", "/documents"),
        ("get", f"/conversions/{uuid.uuid4()}"),
        ("put", f"/conversions/{uuid.uuid4()}"),
        ("get", f"/conversions/{uuid.uuid4()}/export?format=docx"),
    ],
)
def test_인증_없이는_401(client: TestClient, method: str, path: str) -> None:
    response = client.request(method, path, json={"text": _TEXT, "edited_text": _TEXT})

    assert response.status_code == 401
    assert response.headers["WWW-Authenticate"] == "Bearer"
